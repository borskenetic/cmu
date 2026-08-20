"""Generate CMU progress report from template."""
from pathlib import Path

from docx import Document

TEMPLATE = Path(r"c:\Users\THIS PC\Desktop\Reports\DOCS\Project_Progress_Report_2026-08-11.docx")
OUTPUT = Path(r"c:\Users\THIS PC\Desktop\Reports\DOCS\Project_Progress_Report_2026-08-17.docx")

ROWS = [
    (
        "School Setup — College & Course Hierarchy",
        "Renamed Prospectus Manager to School Setup. Added colleges table, College model, and "
        "program.college_id so courses (e.g. Bachelor of Science in Information Technology) nest "
        "under college departments (e.g. College of Education). Rebuilt ProspectusController, "
        "School Setup UI, program_item partial, routes, and seeders with default CMU colleges.",
    ),
    (
        "Attendance Logs — College Filter",
        "Replaced the All Courses dropdown on Attendance Logs with a College filter whose options "
        "are sourced from School Setup. Updated AttendanceLogController filtering and the "
        "attendance_logs index view.",
    ),
    (
        "School Setup UI & Interaction Fixes",
        "Fixed expand/collapse toggle visibility, college click handling, and modal/toast z-index "
        "so Add Course inputs appear and Edit/Delete dialogs open above the sidebar. Refined "
        "prospectus.js and School Setup Blade templates after user-reported freeze and no-click issues.",
    ),
    (
        "Production Deploy — Sidebar CSS 404",
        "Resolved live console 404 for admin-sidebar.css by loading sidebar styles inline only in "
        "layouts/app.blade.php, avoiding a missing public CSS file on the web host while keeping "
        "sidebar styling intact.",
    ),
]


def main() -> None:
    doc = Document(str(TEMPLATE))

    doc.paragraphs[2].text = "Date: 17 August 2026  ·  Client: CMU"

    table = doc.tables[0]

    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._element)

    for index, (module, description) in enumerate(ROWS, start=1):
        row = table.add_row()
        row.cells[0].text = str(index)
        row.cells[1].text = module
        row.cells[2].text = description
        row.cells[3].text = "✓  Completed"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
